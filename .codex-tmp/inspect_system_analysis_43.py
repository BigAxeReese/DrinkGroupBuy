from pathlib import Path
import sys

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.oxml.ns import qn

sys.stdout.reconfigure(encoding="utf-8")

path = Path(r"C:\vscode\DrinkGroupBuy\system-analysis\系統分析書.docx")
doc = Document(path)
extract_dir = Path(r"C:\vscode\DrinkGroupBuy\.codex-tmp\system-analysis-review-20260808\images")
extract_dir.mkdir(parents=True, exist_ok=True)

print(f"DOCUMENT paragraphs={len(doc.paragraphs)} tables={len(doc.tables)} images={len(doc.inline_shapes)}")

table_map = {id(t._tbl): i for i, t in enumerate(doc.tables)}
paragraph_map = {id(p._p): i for i, p in enumerate(doc.paragraphs)}

active = False
for child in doc.element.body.iterchildren():
    if child.tag == qn("w:p"):
        p = Paragraph(child, doc)
        text = p.text.strip().replace("\n", " / ")
        if text.startswith("4.3 ") or text.startswith("4.3顧客"):
            active = True
        if active and (text.startswith("4.4 ") or text.startswith("4.4商家")):
            break
        if active and text:
            index = paragraph_map.get(id(child), -1)
            print(f"P{index} [{p.style.name}] {text}")
    elif child.tag == qn("w:tbl") and active:
        table = Table(child, doc)
        index = table_map.get(id(child), -1)
        print(f"TABLE {index} rows={len(table.rows)} cols={len(table.columns)}")
        for row in table.rows:
            print("  " + " | ".join(cell.text.replace("\n", " / ") for cell in row.cells))

print("ALL 4.3 HEADINGS")
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if text.startswith("4.3") or text.startswith("4.2.2付款"):
        print(f"P{i} [{p.style.name}] {text}")

print("4.3 DRAWINGS")
active = False
for child in doc.element.body.iterchildren():
    if child.tag != qn("w:p"):
        continue
    p = Paragraph(child, doc)
    text = p.text.strip()
    if text.startswith("4.3 ") or text.startswith("4.3顧客"):
        active = True
    if active and (text.startswith("4.4 ") or text.startswith("4.4商家")):
        break
    if not active:
        continue
    blips = child.xpath(".//a:blip")
    for blip in blips:
        rid = blip.get(qn("r:embed"))
        part = doc.part.related_parts[rid]
        print(f"IMAGE paragraph={text!r} rid={rid} part={part.partname} bytes={len(part.blob)}")
        if part.partname.filename in {f"image{i}.png" for i in range(10, 16)}:
            (extract_dir / part.partname.filename).write_bytes(part.blob)
